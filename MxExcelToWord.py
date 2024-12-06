from docx import Document

# 打开现有文档
doc = Document('E:/万合结构/1项目/WHJ82蒸发波导诊断系统/蒸发波导诊断系统框图样目录.docx')

template_section = doc.sections[2]

template_doc = Document()

# 创建一个新的文档
template_doc = Document()

for paragraph in doc.paragraphs:
    template_doc.add_paragraph(paragraph.text)

# # 复制表格
# for table in template_section._element.body.iter_tables():# 表格
#     new_table = template_doc.add_table(rows=0, cols=len(table.columns))
#     for row in table.rows:
#         new_row = new_table.add_row().cells
#         for i, cell in enumerate(row.cells):
#             new_row[i].text = cell.text


template_doc.save('E:/万合结构/1项目/WHJ82蒸发波导诊断系统/template.docx')
